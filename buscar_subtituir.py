from docx import Document


def _substituir_no_paragrafo(paragrafo, pares):
    """
    pares: dict {texto_original: texto_novo}
    Substitui direto no texto do parágrafo, sem depender de marcadores.
    """
    if not paragrafo.runs:
        return

    texto_completo = "".join(run.text for run in paragrafo.runs)

    if not texto_completo:
        return

    novo_texto = texto_completo

    # Troca primeiro os textos mais longos, para evitar que um termo curto
    # (ex: "Ana") acabe mexendo dentro de um termo maior (ex: "Ana Paula").
    for original in sorted(pares.keys(), key=len, reverse=True):
        novo = pares[original]
        if original and original in novo_texto:
            novo_texto = novo_texto.replace(original, novo)

    if novo_texto == texto_completo:
        return

    # Consolida o texto no primeiro run (mantendo a formatação dele)
    # e esvazia os demais runs do parágrafo — necessário porque o Word
    # frequentemente divide o mesmo texto visual em vários runs.
    paragrafo.runs[0].text = novo_texto
    for run in paragrafo.runs[1:]:
        run.text = ""


def contar_ocorrencias(caminho, termo):
    """Conta quantas vezes um termo aparece no documento (parágrafos e tabelas)."""
    doc = Document(caminho)
    total = 0

    for paragrafo in doc.paragraphs:
        total += paragrafo.text.count(termo)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    total += paragrafo.text.count(termo)

    return total


def aplicar_substituicoes(modelo, pares, saida):
    """
    modelo: caminho do docx original
    pares: dict {texto_original: texto_novo}
    saida: caminho do docx gerado
    """
    doc = Document(modelo)

    for paragrafo in doc.paragraphs:
        _substituir_no_paragrafo(paragrafo, pares)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_no_paragrafo(paragrafo, pares)

    doc.save(saida)