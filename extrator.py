import re
from docx import Document

# Aceita {{nome}}, {{ nome }}, {{nome_1}}, etc.
PADRAO = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _extrair_de_paragrafos(paragrafos, vistos, campos):
    for paragrafo in paragrafos:
        # Junta o texto de todos os runs do parágrafo: o Word costuma
        # fragmentar o texto em vários runs, então buscar run a run
        # faz o padrão {{...}} não ser encontrado.
        texto = paragrafo.text

        for match in PADRAO.finditer(texto):
            nome_campo = match.group(1).strip()

            if nome_campo not in vistos:
                vistos.add(nome_campo)
                campos.append({
                    "campo": nome_campo,
                    "texto_original": match.group(0)
                })


def extrair_campos(caminho):
    doc = Document(caminho)

    campos = []
    vistos = set()

    _extrair_de_paragrafos(doc.paragraphs, vistos, campos)

    # Também procura dentro de tabelas, comum em petições (cabeçalhos, qualificação)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                _extrair_de_paragrafos(celula.paragraphs, vistos, campos)

    return campos