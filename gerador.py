import re
from docx import Document

PADRAO = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _substituir_no_paragrafo(paragrafo, dados):
    if not paragrafo.runs:
        return

    texto_completo = "".join(run.text for run in paragrafo.runs)

    if not PADRAO.search(texto_completo):
        return

    def repl(match):
        chave = match.group(1).strip()
        # Se por algum motivo o campo não veio no dicionário, mantém o
        # marcador original em vez de apagar o texto silenciosamente.
        return dados.get(chave, match.group(0))

    novo_texto = PADRAO.sub(repl, texto_completo)

    # Coloca o texto final no primeiro run (que mantém a formatação
    # original daquele trecho) e esvazia os demais runs do parágrafo.
    paragrafo.runs[0].text = novo_texto
    for run in paragrafo.runs[1:]:
        run.text = ""


def gerar_documento(modelo, dados, saida):
    doc = Document(modelo)

    for paragrafo in doc.paragraphs:
        _substituir_no_paragrafo(paragrafo, dados)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    _substituir_no_paragrafo(paragrafo, dados)

    doc.save(saida)